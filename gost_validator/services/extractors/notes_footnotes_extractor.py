"""Извлечение признаков примечаний и сносок."""

from __future__ import annotations

from lxml import etree
from docx import Document
from docx.oxml.ns import qn

from ...config.regex_patterns import (
    RE_ASTERISK_FOOTNOTE_BODY,
    RE_ASTERISK_FOOTNOTE_INLINE_MARKER,
    RE_FIGURE_CAPTION,
    RE_NOTES_HEADER,
    RE_NOTES_ITEM,
    RE_NOTE_SINGLE,
    RE_NUMBERED_NOTE_SINGLE,
    RE_TABLE_CONTINUATION,
    RE_TABLE_TITLE,
)
from ...models.rich_document_structure import FootnoteFeature, NoteFeature
from .common import clean_text


def _near_material_flags(
    paragraph_index: int,
    figure_caption_indices: set[int],
    table_caption_indices: set[int],
    window: int = 2,
) -> tuple[bool, bool]:
    near_figure = any(abs(paragraph_index - idx) <= window for idx in figure_caption_indices)
    near_table = any(abs(paragraph_index - idx) <= window for idx in table_caption_indices)
    return near_figure, near_table


def _footnotes_root(doc: Document) -> object | None:
    for part in doc.part.package.parts:
        part_name = str(getattr(part, "partname", ""))
        if not part_name.endswith("/footnotes.xml"):
            continue

        root = getattr(part, "_element", None)
        if root is None:
            blob = getattr(part, "blob", None)
            if not blob:
                return None
            root = etree.fromstring(blob)
        return root

    return None


def _extract_footnote_ids_from_part(doc: Document) -> set[int]:
    """Возвращает набор id из /word/footnotes.xml, если часть присутствует."""
    root = _footnotes_root(doc)
    if root is None:
        return set()

    result: set[int] = set()
    for node in root.findall(".//w:footnote", root.nsmap):
        raw_id = node.get(qn("w:id"))
        if raw_id is None:
            continue
        try:
            result.add(int(raw_id))
        except ValueError:
            continue
    return result


def _extract_footnote_marker_from_content(doc: Document, footnote_id: int) -> str | None:
    """
    Читает первый символ/слово из содержимого сноски в footnotes.xml.
    Возвращает '*' если сноска начинается с *, иначе None (используем ID как маркер).
    """
    root = _footnotes_root(doc)
    if root is None:
        return None

    for node in root.findall(".//w:footnote", root.nsmap):
        raw_id = node.get(qn("w:id"))
        if raw_id is None:
            continue
        try:
            if int(raw_id) != footnote_id:
                continue
        except ValueError:
            continue

        # Читаем текст сноски
        text_parts = []
        for t_elem in node.findall(".//w:t", root.nsmap):
            if t_elem.text:
                text_parts.append(t_elem.text)

        footnote_text = "".join(text_parts).strip()
        if footnote_text.startswith("*"):
            return "*"
        break

    return None


def _extract_footnote_separator_flags(doc: Document) -> tuple[bool | None, bool | None]:
    """Возвращает признаки разделительной линии для сносок из footnotes.xml."""
    root = _footnotes_root(doc)
    if root is None:
        return None, None

    separator_node = None
    for node in root.findall(".//w:footnote", root.nsmap):
        node_type = (node.get(qn("w:type")) or "").lower()
        if node_type == "separator":
            separator_node = node
            break

    if separator_node is None:
        return False, None

    has_separator = bool(
        separator_node.findall(".//w:separator", separator_node.nsmap)
        or separator_node.findall(".//w:pBdr/w:top", separator_node.nsmap)
    )
    # Для стандартного Word separator считаем эвристику "короткая линия слева" положительной.
    short_left_heuristic = has_separator
    return has_separator, short_left_heuristic


def extract_notes_features(doc: Document) -> list[NoteFeature]:
    """Извлекает абзацы-примечания по шаблонам ТЗ/ГОСТ."""
    notes: list[NoteFeature] = []
    in_notes_group = False

    normalized_paragraphs = [clean_text(paragraph.text) for paragraph in doc.paragraphs]
    figure_caption_indices = {
        index for index, text in enumerate(normalized_paragraphs) if text and RE_FIGURE_CAPTION.match(text)
    }
    table_caption_indices = {
        index
        for index, text in enumerate(normalized_paragraphs)
        if text and (RE_TABLE_TITLE.match(text) or RE_TABLE_CONTINUATION.search(text))
    }

    for paragraph_index, paragraph in enumerate(doc.paragraphs):
        text = clean_text(paragraph.text)
        if not text:
            in_notes_group = False
            continue

        if RE_NOTE_SINGLE.match(text):
            near_figure, near_table = _near_material_flags(paragraph_index, figure_caption_indices, table_caption_indices)
            notes.append(
                NoteFeature(
                    paragraph_index=paragraph_index,
                    raw_text=text,
                    note_kind="single",
                    has_dash_separator=True,
                    near_figure_caption=near_figure,
                    near_table_caption=near_table,
                )
            )
            in_notes_group = False
            continue

        numbered_single_match = RE_NUMBERED_NOTE_SINGLE.match(text)
        if numbered_single_match:
            near_figure, near_table = _near_material_flags(paragraph_index, figure_caption_indices, table_caption_indices)
            notes.append(
                NoteFeature(
                    paragraph_index=paragraph_index,
                    raw_text=text,
                    note_kind="numbered_single",
                    item_number=int(numbered_single_match.group(1)),
                    has_dash_separator=True,
                    near_figure_caption=near_figure,
                    near_table_caption=near_table,
                )
            )
            in_notes_group = False
            continue

        if RE_NOTES_HEADER.match(text):
            near_figure, near_table = _near_material_flags(paragraph_index, figure_caption_indices, table_caption_indices)
            notes.append(
                NoteFeature(
                    paragraph_index=paragraph_index,
                    raw_text=text,
                    note_kind="group_header",
                    has_dash_separator=False,
                    near_figure_caption=near_figure,
                    near_table_caption=near_table,
                )
            )
            in_notes_group = True
            continue

        if in_notes_group:
            group_item_match = RE_NOTES_ITEM.match(text)
            if group_item_match:
                near_figure, near_table = _near_material_flags(paragraph_index, figure_caption_indices, table_caption_indices)
                notes.append(
                    NoteFeature(
                        paragraph_index=paragraph_index,
                        raw_text=text,
                        note_kind="group_item",
                        item_number=int(group_item_match.group(1)),
                        has_dash_separator=None,
                        near_figure_caption=near_figure,
                        near_table_caption=near_table,
                    )
                )
                continue
            in_notes_group = False

    return notes


def extract_footnote_features(doc: Document) -> list[FootnoteFeature]:
    """Извлекает маркеры сносок из body и связывает с /word/footnotes.xml."""
    footnote_features: list[FootnoteFeature] = []
    existing_footnote_ids = _extract_footnote_ids_from_part(doc)
    has_separator_line, separator_short_left = _extract_footnote_separator_flags(doc)

    paragraphs = list(doc.paragraphs)
    asterisk_body_count = sum(1 for paragraph in paragraphs if RE_ASTERISK_FOOTNOTE_BODY.match(paragraph.text or ""))
    inline_asterisk_markers: list[int] = []

    for index, paragraph in enumerate(paragraphs):
        text = paragraph.text or ""
        inline_asterisk_markers.extend([index for _ in RE_ASTERISK_FOOTNOTE_INLINE_MARKER.finditer(text)])

    has_asterisk_resolution = len(inline_asterisk_markers) > 0 and asterisk_body_count > 0

    for paragraph_index, paragraph in enumerate(paragraphs):
        paragraph_element = paragraph._p

        for reference in paragraph_element.findall(".//w:footnoteReference", paragraph_element.nsmap):
            raw_id = reference.get(qn("w:id"))
            if raw_id is None:
                continue

            footnote_id: int | None = None
            try:
                footnote_id = int(raw_id)
            except ValueError:
                pass

            # Читаем маркер из содержимого сноски (если он начинается с *)
            custom_marker = _extract_footnote_marker_from_content(doc, footnote_id) if footnote_id is not None else None
            is_custom_marker = custom_marker == "*"
            marker_text = "*" if is_custom_marker else str(raw_id)

            footnote_features.append(
                FootnoteFeature(
                    paragraph_index=paragraph_index,
                    marker_text=marker_text,
                    marker_type="xml_reference",
                    footnote_id=footnote_id,
                    custom_mark_follows=is_custom_marker,
                    resolved_in_footnotes_part=(footnote_id in existing_footnote_ids) if footnote_id is not None else None,
                    has_separator_line=has_separator_line,
                    separator_short_left_heuristic=separator_short_left,
                )
            )

        text = paragraph.text or ""
        for _ in RE_ASTERISK_FOOTNOTE_INLINE_MARKER.finditer(text):
            footnote_features.append(
                FootnoteFeature(
                    paragraph_index=paragraph_index,
                    marker_text="*",
                    marker_type="asterisk",
                    footnote_id=None,
                    custom_mark_follows=None,
                    resolved_in_footnotes_part=has_asterisk_resolution,
                    has_separator_line=has_separator_line,
                    separator_short_left_heuristic=separator_short_left,
                )
            )

    return footnote_features
