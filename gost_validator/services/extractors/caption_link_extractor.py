"""Извлечение подписей рисунков и ссылок из текста документа."""

from __future__ import annotations

import re

from docx import Document

from ...config.regex_patterns import (
    RE_APPENDIX_HEADER,
    RE_FIGURE_CAPTION,
    RE_FIGURE_LINK,
    RE_FORMULA_LINK,
    RE_SOURCE_LINK,
    RE_TABLE_CONTINUATION,
    RE_TABLE_LINK,
    RE_TABLE_TITLE,
)
from ...models.rich_document_structure import FigureCaptionFeature, LinkFeature
from .common import resolve_paragraph_alignment


def _is_range(target: str | None) -> bool:
    if not target:
        return False
    return bool(re.search(r"[-\u2013]", target))


def _extract_references_numbers(doc: Document) -> set[int] | None:
    """Возвращает номера источников из раздела "Список использованных источников".

    Если раздел не найден, возвращает None.
    """
    paragraphs = [" ".join(p.text.split()) for p in doc.paragraphs]

    start_index = None
    for idx, text in enumerate(paragraphs):
        upper = text.upper()
        if "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ" in upper:
            start_index = idx
            break

    if start_index is None:
        return None

    numbers: set[int] = set()
    numbered_line_re = re.compile(r"^(\d{1,4})[\.)]?\s+.+")

    for text in paragraphs[start_index + 1 :]:
        if not text:
            continue

        # Доходим до приложений и выходим из секции источников.
        if RE_APPENDIX_HEADER.match(text):
            break

        match = numbered_line_re.match(text)
        if match:
            numbers.add(int(match.group(1)))

    return numbers


def _resolve_source_link(start_number: str, end_number: str | None, refs: set[int] | None) -> bool | None:
    if refs is None:
        return None

    try:
        start = int(start_number)
    except ValueError:
        return False

    if end_number is None:
        return start in refs

    try:
        end = int(end_number)
    except ValueError:
        return False

    if end < start:
        return False

    return all(num in refs for num in range(start, end + 1))


def _figure_caption_pattern(number_token: str) -> str:
    if re.match(r"^[A-Za-zА-Яа-я]\.\d+$", number_token):
        return "figure_caption_appendix"
    if re.match(r"^\d+\.\d+$", number_token):
        return "figure_caption_sectional"
    return "figure_caption_global"


def extract_figure_caption_features(doc: Document) -> list[FigureCaptionFeature]:
    """Извлекает подписи рисунков по regex-эвристике."""
    result: list[FigureCaptionFeature] = []
    is_in_appendix_context = False

    for paragraph_index, paragraph in enumerate(doc.paragraphs):
        text = " ".join(paragraph.text.split())
        if not text:
            continue

        # Якорь приложения по ТЗ: после заголовка "ПРИЛОЖЕНИЕ ..." считаем контекст приложением.
        if RE_APPENDIX_HEADER.match(text):
            is_in_appendix_context = True

        match = RE_FIGURE_CAPTION.search(text)
        if not match:
            continue

        number_token = match.group(2)
        separator = match.group(3) or ""
        result.append(
            FigureCaptionFeature(
                paragraph_index=paragraph_index,
                caption_text=text,
                alignment=resolve_paragraph_alignment(paragraph),
                pattern_type=_figure_caption_pattern(number_token),
                has_dash_separator=separator in {"-", "–", "—", ":"},
                ends_with_period=text.endswith("."),
                in_appendix=is_in_appendix_context,
            )
        )

    return result


def extract_links_features(doc: Document) -> list[LinkFeature]:
    """Извлекает ссылки на источники/рисунки/таблицы/формулы."""
    links: list[LinkFeature] = []
    refs_numbers = _extract_references_numbers(doc)

    for paragraph in doc.paragraphs:
        text = " ".join(paragraph.text.split())
        if not text:
            continue

        is_caption_or_title = bool(
            RE_FIGURE_CAPTION.match(text)
            or RE_TABLE_TITLE.match(text)
            or RE_TABLE_CONTINUATION.search(text)
        )

        for match in RE_SOURCE_LINK.finditer(text):
            start_number = match.group("start")
            end_number = match.group("end")
            target = f"{start_number}-{end_number}" if end_number else start_number
            links.append(
                LinkFeature(
                    link_type="source",
                    raw_text=match.group(0),
                    target_number=target,
                    is_range=end_number is not None or _is_range(target),
                    resolved_in_target_list=_resolve_source_link(start_number, end_number, refs_numbers),
                )
            )

        if is_caption_or_title:
            continue

        for match in RE_FIGURE_LINK.finditer(text):
            target = match.group(1)
            links.append(LinkFeature(link_type="figure", raw_text=match.group(0), target_number=target))

        for match in RE_TABLE_LINK.finditer(text):
            target = match.group(1)
            links.append(LinkFeature(link_type="table", raw_text=match.group(0), target_number=target))

        for match in RE_FORMULA_LINK.finditer(text):
            target = match.group(1)
            links.append(LinkFeature(link_type="formula", raw_text=match.group(0), target_number=target))

    return links
