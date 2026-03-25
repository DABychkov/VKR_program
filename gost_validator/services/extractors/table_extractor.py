"""Извлечение признаков таблиц.

Текущий файл — скелет. Реализацию добавляем поэтапно.
"""

from __future__ import annotations

from docx import Document
from docx.oxml.ns import qn

from ...config.regex_patterns import RE_TABLE_CONTINUATION, RE_TABLE_TITLE
from ...models.rich_document_structure import RunFeature, TableCellFeature, TableFeature
from .common import clean_text, extract_run_feature, resolve_paragraph_alignment


def _collect_cell_runs(cell: object) -> list[RunFeature]:
    runs: list[RunFeature] = []
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            runs.append(extract_run_feature(run))
    return runs


def _title_pattern_type(title_text: str | None) -> str | None:
    if not title_text:
        return None
    if RE_TABLE_TITLE.search(title_text):
        return "table_numbered"
    return None


def _is_valid_table_title_text(text: str | None) -> bool:
    if not text:
        return False
    return bool(RE_TABLE_TITLE.search(text) or RE_TABLE_CONTINUATION.search(text))


def _continuation_marker(title_text: str | None) -> str | None:
    if not title_text:
        return None
    if RE_TABLE_CONTINUATION.search(title_text):
        return title_text
    return None


def _collect_table_cell_map(table: object) -> list[TableCellFeature]:
    cells: list[TableCellFeature] = []
    for row_index, row in enumerate(table.rows):
        for col_index, cell in enumerate(row.cells):
            cells.append(
                TableCellFeature(
                    row=row_index,
                    col=col_index,
                    text=clean_text(cell.text),
                    is_header_row=row_index == 0,
                    runs_features=_collect_cell_runs(cell),
                )
            )
    return cells


def _table_inside_borders(table: object) -> tuple[bool | None, bool | None]:
    tbl = table._tbl
    tbl_pr = tbl.find(qn("w:tblPr"))
    if tbl_pr is None:
        return None, None

    tbl_borders = tbl_pr.find(qn("w:tblBorders"))
    if tbl_borders is None:
        return None, None

    inside_h = tbl_borders.find(qn("w:insideH"))
    inside_v = tbl_borders.find(qn("w:insideV"))

    def _exists_border(border_element: object | None) -> bool | None:
        if border_element is None:
            return None
        val = (border_element.get(qn("w:val")) or "").lower()
        if val in {"", "nil", "none"}:
            return False
        return True

    return _exists_border(inside_h), _exists_border(inside_v)


def _table_outer_borders(table: object) -> tuple[bool | None, bool | None, bool | None, bool | None]:
    tbl = table._tbl
    tbl_pr = tbl.find(qn("w:tblPr"))
    if tbl_pr is None:
        return None, None, None, None

    tbl_borders = tbl_pr.find(qn("w:tblBorders"))
    if tbl_borders is None:
        return None, None, None, None

    top = tbl_borders.find(qn("w:top"))
    bottom = tbl_borders.find(qn("w:bottom"))
    left = tbl_borders.find(qn("w:left"))
    right = tbl_borders.find(qn("w:right"))

    def _exists_border(border_element: object | None) -> bool | None:
        if border_element is None:
            return None
        val = (border_element.get(qn("w:val")) or "").lower()
        if val in {"", "nil", "none"}:
            return False
        return True

    return _exists_border(top), _exists_border(bottom), _exists_border(left), _exists_border(right)


def _table_has_diagonal_borders(table: object) -> bool | None:
    found_cells = False
    has_diagonal = False

    for cell in table._tbl.findall(".//w:tc", table._tbl.nsmap):
        found_cells = True
        tc_pr = cell.find(qn("w:tcPr"))
        if tc_pr is None:
            continue
        tc_borders = tc_pr.find(qn("w:tcBorders"))
        if tc_borders is None:
            continue

        for diag_name in ("w:tl2br", "w:tr2bl", "w:diagUp", "w:diagDown"):
            diag = tc_borders.find(qn(diag_name))
            if diag is None:
                continue
            val = (diag.get(qn("w:val")) or "").lower()
            if val not in {"", "nil", "none"}:
                has_diagonal = True
                break

        if has_diagonal:
            break

    if not found_cells:
        return None
    return has_diagonal


def _table_prev_paragraph_map(doc: Document) -> dict[int, int | None]:
    prev_map: dict[int, int | None] = {}
    paragraph_index = 0
    table_index = 0

    for child in doc._element.body.iterchildren():
        if child.tag == qn("w:p"):
            paragraph_index += 1
            continue
        if child.tag == qn("w:tbl"):
            prev_map[table_index] = paragraph_index - 1 if paragraph_index > 0 else None
            table_index += 1

    return prev_map


def extract_table_features(doc: Document) -> list[TableFeature]:
    """Возвращает признаки таблиц документа.

    Этап 2 (минимум):
    - rows_count / cols_count
    - cell_text_map
    - title_above_text (эвристика по предыдущему абзацу)

    TODO:
    - Добавить анализ границ/диагоналей через XML.
    - Добавить распознавание "Продолжение таблицы N".
    """
    table_features: list[TableFeature] = []
    prev_paragraph_map = _table_prev_paragraph_map(doc)

    for table_index, table in enumerate(doc.tables):
        rows_count = len(table.rows)
        cols_count = len(table.columns) if rows_count > 0 else 0
        prev_paragraph_index = prev_paragraph_map.get(table_index)
        prev_paragraph = doc.paragraphs[prev_paragraph_index] if prev_paragraph_index is not None else None

        title_text = None
        title_paragraph_index = None
        title_alignment = "unknown"
        if prev_paragraph is not None:
            candidate = clean_text(prev_paragraph.text)
            if _is_valid_table_title_text(candidate):
                title_text = candidate
                title_paragraph_index = prev_paragraph_index
                title_alignment = resolve_paragraph_alignment(prev_paragraph)

        inside_h, inside_v = _table_inside_borders(table)
        outer_top, outer_bottom, outer_left, outer_right = _table_outer_borders(table)
        cell_map = _collect_table_cell_map(table)
        table_features.append(
            TableFeature(
                table_index=table_index,
                table_anchor_paragraph_index=prev_paragraph_index,
                rows_count=rows_count,
                cols_count=cols_count,
                title_above_text=title_text,
                title_paragraph_index=title_paragraph_index,
                title_alignment=title_alignment,
                title_pattern_type=_title_pattern_type(title_text),
                has_inside_horizontal_borders=inside_h,
                has_inside_vertical_borders=inside_v,
                has_outer_top_border=outer_top,
                has_outer_bottom_border=outer_bottom,
                has_outer_left_border=outer_left,
                has_outer_right_border=outer_right,
                has_diagonal_borders=_table_has_diagonal_borders(table),
                continuation_marker=_continuation_marker(title_text),
                header_row_cells=[cell for cell in cell_map if cell.is_header_row],
                first_column_cells=[cell for cell in cell_map if cell.col == 0],
                cell_text_map=cell_map,
            )
        )

    return table_features
